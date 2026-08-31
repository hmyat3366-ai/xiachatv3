import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_document(output_path):
    doc = Document()

    # Page setup - Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Colors
    primary_orange = RGBColor(220, 100, 20)     # #DC6414
    dark_heading = RGBColor(23, 23, 23)        # #171717
    body_color = RGBColor(50, 50, 50)          # #323232
    gray_subtext = RGBColor(100, 100, 100)     # #646464
    table_header_bg = "171717"
    table_alt_bg = "F9FAFB"
    callout_bg = "FFF7ED"
    callout_border = "FDBA74"

    # Default Style Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Pyidaungsu'
    font.size = Pt(10.5)
    font.color.rgb = body_color

    # TITLE BANNER TABLE
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_cell = title_table.cell(0, 0)
    title_cell.width = Inches(6.9)
    set_cell_background(title_cell, "171717")
    set_cell_margins(title_cell, top=200, bottom=200, left=250, right=250)

    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("Xia Chat — စနစ်တည်ဆောက်မှု လိုအပ်ချက်နှင့် စတင်ရေးသားရန် လမ်းညွှန်\n")
    run_title.font.name = 'Pyidaungsu'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)

    run_subtitle = p_title.add_run("Full Architecture Requirements, Tech Stack & Step-by-Step Development Roadmap")
    run_subtitle.font.name = 'Pyidaungsu'
    run_subtitle.font.size = Pt(11)
    run_subtitle.font.color.rgb = RGBColor(255, 138, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # CALLOUT BOX: Summary
    callout_table = doc.add_table(rows=1, cols=1)
    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_cell = callout_table.cell(0, 0)
    c_cell.width = Inches(6.9)
    set_cell_background(c_cell, callout_bg)
    set_cell_margins(c_cell, top=120, bottom=120, left=180, right=180)
    cp = c_cell.paragraphs[0]
    c_run = cp.add_run("📌 စနစ်အကျဉ်းချုပ် (Executive Summary):\n")
    c_run.font.bold = True
    c_run.font.color.rgb = primary_orange
    c_desc = cp.add_run(
        "Xia Chat သည် လုပ်ငန်းသုံး Omnichannel Customer Support & Unified Inbox SaaS Platform တစ်ခုဖြစ်ပါသည်။ "
        "Website Live Chat, Facebook Messenger, WhatsApp စသည့် ချန်နယ်စုံမှ လာရောက်သော Message များကို နေရာတစ်ခုတည်းမှ စီမံခန့်ခွဲနိုင်ပြီး "
        "AI Auto-Reply (RAG Knowledge Base ပါဝင်) နှင့် Human Agent Takeover/Handoff စနစ်များ တွဲဖက်ပါဝင်ပါသည်။ "
        "ဤစာတမ်းတွင် စနစ်စတင်ရေးသားရန် လိုအပ်ချက်များ (Requirements) နှင့် ဘယ်ကစတင်ရေးသားရမည်ဖြစ်သော အဆင့်ဆင့် လမ်းညွှန် (Step-by-Step Roadmap) ကို ပြည့်စုံစွာ ဖော်ပြထားပါသည်။"
    )
    c_desc.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # SECTION 1: လိုအပ်သော နည်းပညာနှင့် Tools များ (System & Tech Stack Requirements)
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("၁။ စနစ်တည်ဆောက်ရန် လိုအပ်ချက်များ (Prerequisites & Tech Stack)")
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = dark_heading
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Xia Chat ပရောဂျက်ကို အစမှအဆုံး တည်ဆောက်ရန်အတွက် အောက်ပါ Development Tools များ၊ Libraries များနှင့် 3rd-Party Services များ လိုအပ်ပါသည်-"
    )

    # Table 1: Tech Stack
    table1 = doc.add_table(rows=1, cols=3)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ကဏ္ဍ (Category)", "အသုံးပြုသည့် နည်းပညာ (Technology)", "အသုံးပြုရသည့် ရည်ရွယ်ချက် (Purpose)"]
    hdr_cells = table1.rows[0].cells
    col_widths = [Inches(1.8), Inches(2.2), Inches(2.9)]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], table_header_bg)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        hdr_cells[i].width = col_widths[i]

    tech_data = [
        ("Runtime & Environment", "Node.js (v20+) & npm", "Backend Server နှင့် Frontend Build Environment"),
        ("Frontend Framework", "React 19 + TypeScript + Vite", "High-performance Single Page App (SPA) UI"),
        ("Frontend Styling", "Tailwind CSS v4 + Framer Motion", "Modern Glassmorphism Design & Smooth Micro-animations"),
        ("UI Icons", "Lucide React", "Dashboard & Inbox Navigation Icons"),
        ("Backend Framework", "Express.js 5 + tsx", "RESTful API Routes, Middleware & Server-Sent Events (SSE)"),
        ("Database Engine", "SQLite (better-sqlite3) / MongoDB", "Fast ACID-compliant Embedded DB with WAL Mode"),
        ("Security & Auth", "jsonwebtoken + bcryptjs + cookie-parser", "HttpOnly Cookie JWT Authentication, Password Hashing"),
        ("Payment Gateway", "Stripe SDK (stripe)", "SaaS Subscriptions, Checkout Sessions & Invoicing"),
        ("Email Service", "Nodemailer", "Password Reset (6-digit code) & Verification Emails"),
        ("External APIs", "Google OAuth 2.0, Meta Developers", "Google One-Tap Login, Messenger & WhatsApp Webhooks"),
    ]

    for row_idx, data in enumerate(tech_data):
        row_cells = table1.add_row().cells
        bg_color = table_alt_bg if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(data):
            row_cells[c_idx].text = text
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            row_cells[c_idx].width = col_widths[c_idx]

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 2: ဘယ်ကစရေးရမလဲ - အဆင့်ဆင့် တည်ဆောက်ခြင်း Roadmap (Step-by-Step Development Roadmap)
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("၂။ ဘယ်ကစတင်ရေးသားရမလဲ (Step-by-Step Development Roadmap)")
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = dark_heading
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "စနစ်တစ်ခုလုံးကို အမှားအယွင်းမရှိ စနစ်တကျ တည်ဆောက်နိုင်ရန် အောက်ပါ အဆင့် ၇ ဆင့်အတိုင်း ရှေ့နောက်စဉ်ဆက်မပြတ် စတင်ရေးသားရပါမည်-"
    )

    phases = [
        {
            "num": "အဆင့် (၁)",
            "title": "Project Setup & Database Schema Architecture (အခြေခံအဆောက်အအုံ ပြင်ဆင်ခြင်း)",
            "steps": [
                "• Vite + React + TypeScript Frontend Initial Setup ပြုလုပ်ခြင်း။",
                "• Express + TypeScript (`tsx`) Backend Server Architecture သတ်မှတ်ခြင်း။",
                "• SQLite Database Schema တည်ဆောက်ခြင်း (users, workspaces, conversations, messages, knowledge_sources, channels, customers, subscriptions, team_conversations စသည့် Tables ၁၈ ခု)။",
                "• Database Connection တွင် High Concurrency ရရှိရန် WAL Mode (`PRAGMA journal_mode = WAL`) သတ်မှတ်ခြင်း။"
            ]
        },
        {
            "num": "အဆင့် (၂)",
            "title": "Authentication, Security & Onboarding Flow (လုံခြုံရေးနှင့် အကောင့်ဖွင့်စနစ်)",
            "steps": [
                "• Bcryptjs ဖြင့် Password Hashing ပြုလုပ်ပြီး Signup / Login API တည်ဆောက်ခြင်း။",
                "• XSS Attack ကာကွယ်ရန် JWT Token ကို HttpOnly Cookie ဖြင့် ထိန်းသိမ်းခြင်း။",
                "• Forgot Password အတွက် 6-digit Code Generation, SHA256 Hashing နှင့် Email ပို့ဆောင်ခြင်း။",
                "• Google OAuth 2.0 One-Tap Sign In / Sign Up Flow ချိတ်ဆက်ခြင်း။",
                "• အကောင့်ဖွင့်ပြီးပါက Workspace အသစ်ဖန်တီးစေသော Multi-step Onboarding Wizard တည်ဆောက်ခြင်း။"
            ]
        },
        {
            "num": "အဆင့် (၃)",
            "title": "Omnichannel Unified Inbox & Real-Time SSE (အဓိက စကားပြောခန်းစနစ်)",
            "steps": [
                "• Multi-channel Conversations & Messages Data Model နှင့် APIs တည်ဆောက်ခြင်း။",
                "• Server-Sent Events (SSE) ဖြင့် Real-Time Message Streaming ပြုလုပ်ခြင်း။",
                "• AI Auto-Reply စနစ်နှင့် Human Agent Takeover / Return-to-AI Handoff Logic ရေးသားခြင်း။",
                "• Customer Context Sidebar (Tags, Notes, Previous Chats) နှင့် AI Suggested Draft Response ထည့်သွင်းခြင်း။"
            ]
        },
        {
            "num": "အဆင့် (၄)",
            "title": "AI Assistant Agents & RAG Knowledge Base Engine (AI စွမ်းရည် တည်ဆောက်ခြင်း)",
            "steps": [
                "• Multi-Agent Management (Agent အသစ်ဖန်တီးခြင်း၊ Custom System Prompt၊ Tone & Style သတ်မှတ်ခြင်း)။",
                "• Knowledge Base Ingestion: Text Manual Entry, FAQ Parser, Document Upload (PDF/TXT), Website URL Ingestion။",
                "• Text Chunking Algorithm (စကားစုများ ပိုင်းခြားခြင်း) နှင့် Semantic Vector Similarity Search RAG Engine တည်ဆောက်ခြင်း။",
                "• Agent Playground UI ဖြင့် Prompt & RAG စွမ်းဆောင်ရည်ကို တိုက်ရိုက် စမ်းသပ်နိုင်စေခြင်း။"
            ]
        },
        {
            "num": "အဆင့် (၅)",
            "title": "Channels & Integrations (ချန်နယ်များ ချိတ်ဆက်ခြင်း)",
            "steps": [
                "• Website Live Chat Widget Embed Code Generator (`<script>` embed code) နှင့် Public Widget API တည်ဆောက်ခြင်း။",
                "• Meta Webhooks (Facebook Messenger, WhatsApp Business) လက်ခံရန် Webhook Verification & Processing API ရေးသားခြင်း။",
                "• Customer Identity Mapping စနစ် (မတူညီသော Channel များမှ Customer ID များကို တစ်စုတစ်စည်းတည်း ချိတ်ဆက်ခြင်း)။"
            ]
        },
        {
            "num": "အဆင့် (၆)",
            "title": "Team Collaboration & Internal Team Chat (အဖွဲ့လိုက် လုပ်ဆောင်မှုစနစ်)",
            "steps": [
                "• Role-Based Access Control - RBAC (Owner, Admin, Member) သတ်မှတ်ခြင်း။",
                "• Workspace Email Invitation Flow (Token-based Invite Links) တည်ဆောက်ခြင်း။",
                "• လုပ်ငန်းအတွင်း ဝန်ထမ်းအချင်းချင်း စကားပြောနိုင်သော Team Chat (1-on-1 Direct Chat & Group Conversations) ထည့်သွင်းခြင်း။",
                "• Team Audit Logs (လုံခြုံရေးမှတ်တမ်းများ) ခြေရာခံခြင်း။"
            ]
        },
        {
            "num": "အဆင့် (၇)",
            "title": "Analytics, Settings & Stripe Billing (အစီရင်ခံစာနှင့် ငွေပေးချေမှုစနစ်)",
            "steps": [
                "• Real Database Metrics Overview (Chat Volume, Resolution Time, CSAT Score, Channel Breakdown)။",
                "• 7-day, 30-day, 90-day, Custom Date Range Filters နှင့် CSV Export စနစ်။",
                "• Stripe Checkout Session & Stripe Customer Portal Integration (Starter, Pro, Enterprise Plans)။",
                "• Plan Limit Middleware (Free/Starter Plan အလိုက် AI Agent အရေအတွက်၊ Knowledge Chunks နှင့် Member Limit များကို ထိန်းချုပ်ခြင်း)။"
            ]
        }
    ]

    for phase in phases:
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(8)
        p_hdr.paragraph_format.space_after = Pt(2)
        r_tag = p_hdr.add_run(f"📌 {phase['num']} — ")
        r_tag.font.bold = True
        r_tag.font.color.rgb = primary_orange
        r_t = p_hdr.add_run(phase['title'])
        r_t.font.bold = True
        r_t.font.color.rgb = dark_heading
        
        for step in phase['steps']:
            p_step = doc.add_paragraph(step)
            p_step.paragraph_format.left_indent = Inches(0.2)
            p_step.paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # SECTION 3: လက်ရှိ Codebase ကို စတင် Run နည်းနှင့် စမ်းသပ်နည်း (How to Run and Test)
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("၃။ လက်ရှိ Project ကို စတင် Run နည်းနှင့် စမ်းသပ်နည်း (Local Execution)")
    h3_run.font.size = Pt(14)
    h3_run.font.bold = True
    h3_run.font.color.rgb = dark_heading
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("လက်ရှိ Codebase ကို Local Computer ပေါ်တွင် စတင် Run ရန် အောက်ပါ အဆင့်များကို လိုက်နာဆောင်ရွက်ပါ-")

    # Run Table
    table2 = doc.add_table(rows=1, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2_headers = ["အဆင့် (Step)", "Run ရမည့် Command / ရှင်းလင်းချက်"]
    t2_cells = table2.rows[0].cells
    t2_widths = [Inches(2.2), Inches(4.7)]
    for i, title in enumerate(t2_headers):
        t2_cells[i].text = title
        t2_cells[i].paragraphs[0].runs[0].font.bold = True
        t2_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(t2_cells[i], table_header_bg)
        set_cell_margins(t2_cells[i], top=100, bottom=100, left=120, right=120)
        t2_cells[i].width = t2_widths[i]

    run_steps = [
        ("၁။ Dependencies သွင်းခြင်း", "npm install\n(ပရောဂျက်အတွက် လိုအပ်သော Packages အားလုံးကို ဒေါင်းလုဒ်ဆွဲသွင်းပါမည်)"),
        ("၂။ Development Server Run ခြင်း", "npm run dev\n(Backend Server: Port 5000 နှင့် Frontend Vite: Port 5173 ကို တစ်ပြိုင်နက် Run ပေးပါသည်)"),
        ("၃။ Frontend သီးသန့် Run ခြင်း", "npm run dev:frontend\n(Vite Frontend UI ကိုသာ သီးသန့် Run လိုပါက အသုံးပြုနိုင်ပါသည်)"),
        ("၄။ Backend Server သီးသန့် Run ခြင်း", "npm run server\n(Node.js Express Backend API ကိုသာ သီးသန့် စမ်းသပ်လိုပါက)"),
        ("၅။ Production Build စစ်ဆေးခြင်း", "npm run build\n(TypeScript type checking နှင့် Vite production bundling ပြုလုပ်ပါသည်)"),
        ("၆။ Code Linter စစ်ဆေးခြင်း", "npm run lint\n(Oxlint ဖြင့် Code အရည်အသွေးနှင့် အမှားများကို စစ်ဆေးပါသည်)")
    ]

    for row_idx, data in enumerate(run_steps):
        row_cells = table2.add_row().cells
        bg_color = table_alt_bg if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(data):
            row_cells[c_idx].text = text
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            row_cells[c_idx].width = t2_widths[c_idx]

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 4: PRODUCTION DEPLOYMENT & ENVIRONMENT VARIABLES (.env)
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("၄။ Production Live လွှင့်ရာတွင် လိုအပ်သော Configuration Keys (.env Checklist)")
    h4_run.font.size = Pt(14)
    h4_run.font.bold = True
    h4_run.font.color.rgb = dark_heading
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Production Server (Vercel, Render, AWS, VPS) ပေါ်သို့ အမှန်တကယ် Live လွှင့်သည့်အခါ `.env` ဖိုင်တွင် အောက်ပါ API Keys များကို ထည့်သွင်းပေးရပါမည်-"
    )

    env_items = [
        ("PORT", "5000", "Backend Server နားထောင်မည့် Port"),
        ("FRONTEND_URL", "https://app.xiachat.ai", "CORS ချိတ်ဆက်မည့် Production Domain URL"),
        ("JWT_SECRET", "super_secret_jwt_key_2026", "Session Token များ Encode/Decode လုပ်ရန် Secret Key"),
        ("STRIPE_SECRET_KEY", "sk_live_...", "Stripe Live API Key (အမှန်တကယ် ငွေလက်ခံရန်)"),
        ("STRIPE_WEBHOOK_SECRET", "whsec_...", "Stripe Webhook Signature Verification Key"),
        ("GOOGLE_CLIENT_ID", "...apps.googleusercontent.com", "Google Cloud Console မှ OAuth Client ID"),
        ("GOOGLE_CLIENT_SECRET", "GOCSPX-...", "Google Cloud Console မှ OAuth Client Secret"),
        ("EMAIL_HOST & EMAIL_PASS", "smtp.sendgrid.net", "Email Verification နှင့် Password Reset အတွက် SMTP Credentials"),
        ("META_APP_SECRET & WHATSAPP_TOKEN", "EAA...", "Facebook Messenger & WhatsApp Business Webhooks"),
    ]

    table3 = doc.add_table(rows=1, cols=3)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3_headers = ["Environment Variable", "ဥပမာတန်ဖိုး (Example)", "အသုံးဝင်ပုံ (Description)"]
    t3_cells = table3.rows[0].cells
    t3_widths = [Inches(2.5), Inches(2.2), Inches(2.2)]
    for i, title in enumerate(t3_headers):
        t3_cells[i].text = title
        t3_cells[i].paragraphs[0].runs[0].font.bold = True
        t3_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(t3_cells[i], table_header_bg)
        set_cell_margins(t3_cells[i], top=100, bottom=100, left=120, right=120)
        t3_cells[i].width = t3_widths[i]

    for row_idx, data in enumerate(env_items):
        row_cells = table3.add_row().cells
        bg_color = table_alt_bg if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(data):
            row_cells[c_idx].text = text
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            row_cells[c_idx].width = t3_widths[c_idx]

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # FOOTER NOTE
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("— Xia Chat Project Architecture & Development Blueprint Document —")
    r_foot.font.size = Pt(9)
    r_foot.font.italic = True
    r_foot.font.color.rgb = gray_subtext

    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    out_dir = os.getcwd()
    file_name = "Xia_Chat_Requirements_and_Roadmap.docx"
    target_path = os.path.join(out_dir, file_name)
    create_document(target_path)
